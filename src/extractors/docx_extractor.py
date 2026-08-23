from pathlib import Path
from docx import Document

def extract_text_from_docx(file_path: str | Path) -> str:
    """
    Extract text from a DOCX document.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        The extracted text as a single string.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the provided file is not a DOCX file.
    """

    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    if path.suffix.lower() != ".docx":
        raise ValueError(
            f"Expected a .docx file, received: {path.suffix}"
        )

    document = Document(str(path))

    paragraphs = []

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()

                if text:
                    paragraphs.append(text)

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            paragraphs.append(text)

    return "\n".join(paragraphs)
